from duckduckgo_search import DDGS
from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypandoc
import pdfplumber
import os
import re
import tempfile

@tool
def web_search(query: str) -> str:
    """Search the web for current information. Use this when the user asks about recent events, current prices, latest versions, news, weather, or anything requiring up-to-date data."""
    try:
        results = DDGS().text(query, max_results=5)
        if not results:
            return "No results found."
        output = []
        for r in results:
            output.append(f"{r['title']}\n{r['body']}\nSource: {r['href']}")
        return "\n\n".join(output)
    except Exception as e:
        return f"Search failed: {str(e)}"

# ── Document ─────────────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def _add_horizontal_rule(doc):
    """Add a thin horizontal line as a paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def _apply_inline_formatting(para, text: str):
    """Handle **bold** and *italic* inline markers in a paragraph."""
    # Combined pattern for bold and italic
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

def _parse_markdown_table(lines: list) -> list[list[str]]:
    """Parse markdown table lines into a 2D list, skipping separator row."""
    rows = []
    for line in lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue  # skip separator row like |---|---|
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

def _add_table(doc, rows: list[list[str]]):
    """Add a properly formatted Word table from a 2D list."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            _apply_inline_formatting(para, cell_text)
            if i == 0:  # header row — bold
                for run in para.runs:
                    run.bold = True
            # cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)

    doc.add_paragraph()  # spacing after table

def _markdown_to_docx(markdown_content: str, output_path: str) -> None:
    doc = Document()

    # Set reasonable margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    lines = markdown_content.splitlines()
    i = 0
    prev_was_blank = False

    while i < len(lines):
        line = lines[i]

        # ── Markdown table ──────────────────────────────────────────
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_markdown_table(table_lines)
            _add_table(doc, rows)
            prev_was_blank = False
            continue

        # ── Horizontal rule ─────────────────────────────────────────
        if line.strip() in ('---', '***', '___'):
            _add_horizontal_rule(doc)
            i += 1
            prev_was_blank = False
            continue

        # ── Blank line — deduplicate ─────────────────────────────────
        if line.strip() == '':
            if not prev_was_blank:
                doc.add_paragraph()
            prev_was_blank = True
            i += 1
            continue

        prev_was_blank = False

        # ── Headings ─────────────────────────────────────────────────
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)

        # ── Bullet list ──────────────────────────────────────────────
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _apply_inline_formatting(para, line[2:].strip())

        # ── Numbered list ────────────────────────────────────────────
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            para = doc.add_paragraph(style='List Number')
            _apply_inline_formatting(para, text.strip())

        # ── Blockquote ───────────────────────────────────────────────
        elif line.startswith('> '):
            para = doc.add_paragraph(style='Quote')
            _apply_inline_formatting(para, line[2:].strip())

        # ── Regular paragraph ────────────────────────────────────────
        else:
            para = doc.add_paragraph()
            _apply_inline_formatting(para, line.strip())

        i += 1

    doc.save(output_path)

def write_document(filename: str, markdown_content: str, cwd: str) -> str:
    """Write a well-formatted .docx file from markdown content."""
    try:
        if not filename.endswith(".docx"):
            filename += ".docx"
        output_path = os.path.join(cwd, filename)
        _markdown_to_docx(markdown_content, output_path)
        return f"Document saved: {output_path}"
    except Exception as e:
        return f"Failed to write document: {str(e)}"

def read_file(filepath: str) -> str:
    """Read and return the text content of a .docx, .pdf, or .txt file."""
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"

        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".docx":
            doc = Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext == ".pdf":
            with pdfplumber.open(filepath) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n\n".join(pages) if pages else "No text found in PDF."

        elif ext in [".txt", ".md"]:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()

        else:
            return f"Unsupported file type: {ext}. Supported: .docx, .pdf, .txt, .md"

    except Exception as e:
        return f"Failed to read file: {str(e)}"